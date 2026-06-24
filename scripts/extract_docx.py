import docx
import sys

def get_text(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.replace('\n', ' '))
            fullText.append(' | '.join(row_data))
    return '\n'.join(fullText)

if __name__ == '__main__':
    text = get_text(sys.argv[1])
    with open(sys.argv[2], 'w', encoding='utf-8') as f:
        f.write(text)
